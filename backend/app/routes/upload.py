from __future__ import annotations
import hashlib
import logging
from pathlib import Path
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from fastapi.responses import JSONResponse
from app.auth.deps import current_user, verify_csrf
from app.config import get_settings
from app.security import safe_upload_path

log = logging.getLogger(__name__)
router = APIRouter(prefix="/api/upload", tags=["upload"])

ALLOWED_EXTENSIONS = {".doc", ".docx", ".pdf"}


@router.post("", dependencies=[Depends(verify_csrf)])
async def upload_file(
    file: UploadFile = File(...),
    user_id: str = Depends(current_user),
) -> JSONResponse:
    s = get_settings()
    # 扩展名校验
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, f"不支持的文件格式: {suffix}，只支持 {ALLOWED_EXTENSIONS}")
    # 文件大小校验（先检查 content-length）
    if file.size and file.size > s.max_upload_mb * 1024 * 1024:
        raise HTTPException(413, f"文件超过 {s.max_upload_mb}MB 上限")
    # 路径穿越防护
    dest = safe_upload_path(s.upload_dir, user_id, file.filename or "unnamed")
    # 读取并保存（带大小二次检查）
    content = await file.read()
    if len(content) > s.max_upload_mb * 1024 * 1024:
        raise HTTPException(413, f"文件超过 {s.max_upload_mb}MB 上限")
    file_hash = hashlib.sha256(content).hexdigest()
    dest.write_bytes(content)
    log.info("upload OK user=%s file=%s size=%dB", user_id, dest.name, len(content))
    return JSONResponse({"file_id": dest.name, "path": str(dest), "size": len(content), "file_hash": file_hash})


@router.get("/preview/{file_name}")
async def preview_file(
    file_name: str,
    user_id: str = Depends(current_user),
):
    from fastapi.responses import FileResponse
    s = get_settings()
    dest = safe_upload_path(s.upload_dir, user_id, file_name)
    if not dest.exists():
        raise HTTPException(404, "文件不存在")
    media = "application/pdf" if dest.suffix.lower() == ".pdf" else "application/octet-stream"
    return FileResponse(str(dest), media_type=media, filename=dest.name)


@router.get("/pdf/{file_name}")
async def preview_as_pdf(
    file_name: str,
    user_id: str = Depends(current_user),
):
    """把 .doc/.docx 用 Word COM 导出为 PDF，浏览器直接渲染（排版与 Word 完全一致）。"""
    from fastapi.responses import FileResponse
    import re
    import sys

    s = get_settings()
    dest = safe_upload_path(s.upload_dir, user_id, file_name)
    if not dest.exists():
        raise HTTPException(404, "文件不存在")

    suffix = dest.suffix.lower()
    if suffix == ".pdf":
        return FileResponse(str(dest), media_type="application/pdf")

    # 缓存目录
    from src.config import get_default_config
    cfg = get_default_config()
    doc_id = re.sub(r"[^A-Za-z0-9_-]+", "_", dest.stem)[:80]
    cache_dir = Path(cfg["paths"]["data_dir"]) / "tmp" / doc_id
    cache_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = cache_dir / (dest.stem + ".pdf")

    # 如果 PDF 缓存已存在直接返回
    if pdf_path.exists():
        return FileResponse(str(pdf_path), media_type="application/pdf")

    # 优先：Windows + 本机 Word（win32com）导出 PDF（在线程池里运行，需 CoInitialize）
    # 回退：LibreOffice headless 转 PDF（适用于无 Word 的环境）
    import asyncio
    def _export_pdf_via_word():
        import win32com.client, pythoncom
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(str(dest.resolve()), ReadOnly=True)
            doc.SaveAs2(str(pdf_path.resolve()), FileFormat=17)  # wdFormatPDF=17
            log.info("word→pdf OK: %s", pdf_path.name)
        finally:
            if doc:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    def _export_pdf_via_libreoffice():
        from src.parse.doc_converter import convert_docx_to_pdf
        converted = convert_docx_to_pdf(dest, cache_dir)
        # convert_docx_to_pdf 会输出到 cache_dir，且文件名为 dest.stem + ".pdf"
        if converted.resolve() != pdf_path.resolve():
            # 保险起见：如果转换器输出了不同路径，复制/重命名到约定缓存路径
            pdf_path.write_bytes(Path(converted).read_bytes())
        log.info("libreoffice→pdf OK: %s", pdf_path.name)

    try:
        loop = asyncio.get_event_loop()
        if sys.platform == "win32":
            try:
                await loop.run_in_executor(None, _export_pdf_via_word)
            except Exception as e:
                log.warning("Word COM 导出失败，回退 LibreOffice: %s", e)
                await loop.run_in_executor(None, _export_pdf_via_libreoffice)
        else:
            await loop.run_in_executor(None, _export_pdf_via_libreoffice)
    except Exception as e:
        log.warning("PDF 导出失败: %s", e)
        raise HTTPException(500, f"PDF 导出失败: {e}")

    return FileResponse(str(pdf_path), media_type="application/pdf")


@router.get("/text/{file_name}")
async def preview_text(
    file_name: str,
    user_id: str = Depends(current_user),
):
    """提取文档纯文本，供前端原文展示。"""
    from fastapi.responses import PlainTextResponse
    import re
    from src.parse.doc_converter import convert_doc_to_docx

    s = get_settings()
    dest = safe_upload_path(s.upload_dir, user_id, file_name)
    if not dest.exists():
        raise HTTPException(404, "文件不存在")

    suffix = dest.suffix.lower()

    try:
        if suffix == ".pdf":
            import pdfplumber
            lines = []
            with pdfplumber.open(str(dest)) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    lines.append(f"{'='*40} 第 {i} 页 {'='*40}")
                    text = page.extract_text() or ""
                    lines.append(text)
            text = "\n".join(lines)

        else:
            # .doc → 找已转换的 .docx 或重新转
            if suffix == ".doc":
                from src.config import get_default_config
                cfg = get_default_config()
                doc_id = re.sub(r"[^A-Za-z0-9_-]+", "_", dest.stem)[:80]
                tmp_dir = Path(cfg["paths"]["data_dir"]) / "tmp" / doc_id
                converted = tmp_dir / (dest.stem + ".docx")
                if not converted.exists():
                    converted = convert_doc_to_docx(dest, tmp_dir)
                docx_path = converted
            else:
                docx_path = dest

            from docx import Document
            doc = Document(str(docx_path))
            lines = []
            for para in doc.paragraphs:
                lines.append(para.text)
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    lines.append("\t".join(c.text for c in row.cells))
            text = "\n".join(lines)

    except Exception as e:
        raise HTTPException(500, f"文本提取失败: {e}")

    return PlainTextResponse(text, media_type="text/plain; charset=utf-8")


@router.get("/html/{file_name}")
async def preview_html(
    file_name: str,
    user_id: str = Depends(current_user),
):
    """把 .doc/.docx 转成 HTML 返回，供前端 iframe srcdoc 渲染。"""
    from fastapi.responses import HTMLResponse
    import mammoth, re, io
    from src.parse.doc_converter import convert_doc_to_docx

    s = get_settings()
    dest = safe_upload_path(s.upload_dir, user_id, file_name)
    if not dest.exists():
        raise HTTPException(404, "文件不存在")

    suffix = dest.suffix.lower()
    if suffix == ".pdf":
        raise HTTPException(400, "PDF 请使用 /preview 接口")

    # .doc → 先找已转换的 .docx，找不到就重新转
    if suffix == ".doc":
        from src.config import get_default_config
        cfg = get_default_config()
        doc_id = re.sub(r"[^A-Za-z0-9_-]+", "_", dest.stem)[:80]
        tmp_dir = Path(cfg["paths"]["data_dir"]) / "tmp" / doc_id
        converted = tmp_dir / (dest.stem + ".docx")
        if not converted.exists():
            try:
                converted = convert_doc_to_docx(dest, tmp_dir)
            except Exception as e:
                raise HTTPException(500, f"文档转换失败: {e}")
        docx_path = converted
    else:
        docx_path = dest

    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(f)

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  body {{ font-family: "Microsoft YaHei", SimSun, sans-serif; font-size: 14px;
         line-height: 1.8; padding: 32px 48px; color: #1e293b; max-width: 860px; margin: 0 auto; }}
  h1,h2,h3,h4 {{ font-weight: 700; margin: 1.2em 0 0.5em; }}
  h1 {{ font-size: 1.5em; }} h2 {{ font-size: 1.3em; }} h3 {{ font-size: 1.1em; }}
  p {{ margin: 0.5em 0; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
  td, th {{ border: 1px solid #d1d5db; padding: 6px 10px; }}
  th {{ background: #f1f5f9; font-weight: 600; }}
  img {{ max-width: 100%; }}
</style>
</head><body>{result.value}</body></html>"""

    return HTMLResponse(html)


@router.get("")
async def list_files(user_id: str = Depends(current_user)) -> JSONResponse:
    s = get_settings()
    user_dir = (s.upload_dir / user_id)
    if not user_dir.exists():
        return JSONResponse({"files": []})
    files = [
        {"name": f.name, "size": f.stat().st_size, "suffix": f.suffix}
        for f in user_dir.iterdir() if f.is_file() and f.suffix.lower() in ALLOWED_EXTENSIONS
    ]
    return JSONResponse({"files": files})
